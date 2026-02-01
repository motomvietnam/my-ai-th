import streamlit as st
import pandas as pd
import re
from io import BytesIO
import docx
import PyPDF2
import difflib
import zipfile
from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import streamlit as st
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
from docx import Document as DocxDocument # Để tránh trùng tên nếu cần

def gop_cac_file_word(list_docx_streams):
    # Tạo một file word mới dựa trên file đầu tiên
    merged_document = DocxDocument(BytesIO(list_docx_streams[0]))
    
    for i in range(1, len(list_docx_streams)):
        # Thêm ngắt trang trước khi nối file tiếp theo
        merged_document.add_page_break()
        
        # Đọc nội dung file tiếp theo
        sub_doc = DocxDocument(BytesIO(list_docx_streams[i]))
        for element in sub_doc.element.body:
            merged_document.element.body.append(element)
            
    output = BytesIO()
    merged_document.save(output)
    return output.getvalue()
def tao_file_word_mau_giay_moi():
    doc = Document()
    # ... code tạo giấy mời ...
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

def tao_file_word_mau_hop_dong():
    doc = Document()
    # ... code tạo hợp đồng ...
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()

def tạo_excel_mẫu():
    # ... code tạo excel ...
    return output.getvalue()

def doc_so_thanh_chu_logic(so_tien):
    # Đảm bảo hàm này cũng nằm ở đây để tránh lỗi NameError lúc xuất ZIP
    return f"{so_tien} đồng"

# --- [DÒNG 101 TRỞ ĐI]: MỚI ĐẾN PHẦN GIAO DIỆN TABS ---
# tabs = st.tabs([...])
# with tabs[2]:
#    ... gọi các hàm ở đây ...
# 1. CẤU HÌNH GIAO DIỆN
st.set_page_config(page_title="Smart Tools Hub - Pro", layout="wide")

st.markdown("""
    <style>
    .stApp { background-color: #f1f5f9; }
    [data-testid="stSidebar"] { background: linear-gradient(180deg, #745af2 0%, #01caf1 100%); }
    [data-testid="stSidebarNav"] ul li div a span { color: white !important; font-size: 18px !important; font-weight: bold !important; }
    div.stButton > button { border-radius: 8px; font-weight: 600; background-color: #745af2; color: white; border: none; width: 100%; }

    [data-testid="stFileUploader"] {
        background-color: #bdc3c7 !important;
        border: 2px dashed #ffffff;
        border-radius: 10px;
        padding: 10px;
    }
    [data-testid="stFileUploader"] section div div { color: white !important; }
    [data-testid="stFileUploader"] svg { fill: white !important; }
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p { font-size: 16px; font-weight: bold; }
    </style>
    """, unsafe_allow_html=True)

if st.sidebar.button("🏠 VỀ DASHBOARD TỔNG"):
    st.switch_page("app.py")

# --- CÁC HÀM HỖ TRỢ ---
def read_file_content(uploaded_file):
    if uploaded_file is None: return ""
    try:
        suffix = uploaded_file.name.split('.')[-1].lower()
        if suffix == 'txt': 
            return str(uploaded_file.read(), "utf-8")
        elif suffix in ['doc', 'docx']:
            doc = docx.Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        elif suffix == 'pdf':
            pdf_reader = PyPDF2.PdfReader(uploaded_file)
            text = "".join([page.extract_text() for page in pdf_reader.pages])
            return text
        elif suffix in ['xlsx', 'xls']:
            df = pd.read_excel(uploaded_file)
            return df.to_string()
    except Exception as e:
        return f"Lỗi đọc file: {e}"
    return ""
def tao_file_word_mau_giay_moi():
    doc = Document()
    
    # Thiết lập Font chữ mặc định
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)

    # 1. Tiêu ngữ
    header = doc.add_paragraph()
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = header.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n")
    run.bold = True
    run = header.add_run("Độc lập - Tự do - Hạnh phúc\n")
    run.bold = True
    header.add_run("----------o0o----------")

    doc.add_paragraph("\n")

    # 2. Tên giấy mời
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("GIẤY MỜI THAM DỰ SỰ KIỆN")
    run.bold = True
    run.font.size = Pt(20)

    doc.add_paragraph("\n")

    # 3. Nội dung mời (Sử dụng các từ khóa {{ }})
    content = doc.add_paragraph()
    content.add_run("Trân trọng kính mời Ông/Bà: ").bold = True
    content.add_run("{{TenKhach}}") # Từ khóa khớp với bảng của bạn
    
    doc.add_paragraph(f"Đến tham dự chương trình: ").add_run("{{TenSuKien}}").bold = True
    
    p = doc.add_paragraph("Thời gian: ")
    p.add_run("{{ThoiGian}}")
    
    p = doc.add_paragraph("Địa điểm: ")
    p.add_run("{{DiaDiem}}")
    
    doc.add_paragraph("\nNội dung sự kiện: Chương trình được tổ chức nhằm tri ân khách hàng và giới thiệu các tính năng mới của công ty.")
    
    # 4. Chữ ký
    doc.add_paragraph("\n")
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("........, ngày .... tháng .... năm 2026\n").italic = True
    run = footer.add_run("ĐẠI DIỆN BAN TỔ CHỨC")
    run.bold = True
    
    # Lưu vào bộ nhớ tạm
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()
def tao_file_word_mau_hop_dong():
    doc = Document()
    
    # 1. Thiết lập Font chữ chuẩn (Times New Roman)
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13) # Thường văn bản chính thức dùng 13pt

    # 2. Tiêu ngữ (Căn giữa, đậm)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM")
    run.bold = True
    run.font.size = Pt(12)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run("Độc lập - Tự do - Hạnh phúc")
    run2.bold = True
    run2.font.size = Pt(13)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.add_run("---------------")

    # 3. Tên hợp đồng
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("\nHỢP ĐỒNG LAO ĐỘNG")
    run_title.bold = True
    run_title.font.size = Pt(16)

    # 4. Nội dung (Sử dụng các biến khớp với bảng Excel của bạn)
    doc.add_paragraph(f"\nChúng tôi, một bên là Công ty: ").add_run("{{TenCongTy}}").bold = True
    doc.add_paragraph(f"Và một bên là Ông/Bà: ").add_run("{{Ten}}").bold = True
    
    # Tạo danh sách thông tin gọn gàng
    fields = [
        ("Mã nhân viên:", "{{MaNV}}"),
        ("Chức vụ:", "{{ChucVu}}"),
        ("Mức lương chính thức:", "{{Luong}}"),
        ("Đơn vị công tác:", "{{Phongban}}"),
        ("Ngày có hiệu lực:", "{{NgayHieuLuc}}")
    ]
    
    for label, placeholder in fields:
        p = doc.add_paragraph()
        p.add_run(f"- {label} ").bold = False
        p.add_run(placeholder).bold = True

    doc.add_paragraph("\nCác điều khoản khác được thực hiện theo quy định của pháp luật lao động hiện hành.")

    # 5. Bảng ký tên (Căn chỉnh chuyên nghiệp)
    doc.add_paragraph("\n")
    table = doc.add_table(rows=1, cols=2)
    table.autofit = True
    
    # Ô bên trái: Người lao động
    cell_left = table.cell(0, 0)
    p_left = cell_left.paragraphs[0]
    p_left.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_l = p_left.add_run("NGƯỜI LAO ĐỘNG")
    run_l.bold = True
    p_left.add_run("\n(Ký và ghi rõ họ tên)")

    # Ô bên phải: Đại diện công ty
    cell_right = table.cell(0, 1)
    p_right = cell_right.paragraphs[0]
    p_right.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_r = p_right.add_run("ĐẠI DIỆN CÔNG TY")
    run_r.bold = True
    p_right.add_run("\n(Ký và đóng dấu)")
    
    # 6. Xuất file
    target_stream = BytesIO()
    doc.save(target_stream)
    return target_stream.getvalue()
def chuan_hoa_excel_pro(df):
    df_clean = df.copy()
    for col in df_clean.columns:
        col_lower = col.lower()
        if any(kw in col_lower for kw in ['tên', 'name', 'họ']):
            df_clean[col] = df_clean[col].apply(lambda x: " ".join(str(x).strip().title().split()) if pd.notnull(x) else x)
        elif any(kw in col_lower for kw in ['sđt', 'đt', 'phone', 'tel']):
            def clean_phone(p):
                n = re.sub(r'\D', '', str(p))
                if n.startswith('84'): n = '0' + n[2:]
                elif not n.startswith('0') and len(n) > 0: n = '0' + n
                return n[-10:] if len(n) > 10 else n
            df_clean[col] = df_clean[col].astype(str).apply(clean_phone)
        elif any(kw in col_lower for kw in ['ngày', 'date']):
            temp_date = pd.to_datetime(df_clean[col], errors='coerce', dayfirst=True)
            df_clean[col] = temp_date.dt.strftime('%d/%m/%Y').fillna('')

    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_clean.to_excel(writer, index=False, sheet_name='Clean_Data')
        workbook = writer.book
        header_fmt = workbook.add_format({'bold': True, 'bg_color': '#745af2', 'font_color': 'white', 'border': 1, 'font_name': 'Arial', 'align': 'center'})
        cell_fmt = workbook.add_format({'border': 1, 'font_name': 'Arial', 'font_size': 11})
        for col_num, value in enumerate(df_clean.columns.values):
            writer.sheets['Clean_Data'].write(0, col_num, value, header_fmt)
            writer.sheets['Clean_Data'].set_column(col_num, col_num, 25, cell_fmt)
    return output.getvalue()

# --- GIAO DIỆN CHÍNH ---
st.title("🚀 SMART TOOLS HUB - EXCEL & DOC PRO")
st.divider()

tabs = st.tabs(["📊 Chuẩn hoá Excel", "🔍 So sánh đối soát", "🎭 Mail Merge (Trộn file)", "💰 Đọc Số Tiền", "📧 Check Email"])

# --- TAB 1: CHUẨN HÓA EXCEL ---
with tabs[0]:
    st.header("📊 Chuẩn hoá Dữ liệu Excel")
    file_ex = st.file_uploader("Kéo thả file Excel tại đây", type=["xlsx"], key="excel_tab")
    if file_ex:
        df = pd.read_excel(file_ex)
        st.dataframe(df.head(10), use_container_width=True)
        if st.button("✨ BẮT ĐẦU CHUẨN HOÁ", key="btn_clean"):
            with st.spinner("Đang xử lý..."):
                res = chuan_hoa_excel_pro(df)
                st.success("✅ Thành công!")
                st.download_button("📥 TẢI FILE EXCEL SẠCH", res, f"Cleaned_{file_ex.name}")

# --- TAB 2: SO SÁNH VĂN BẢN ---
with tabs[1]:
    st.header("🔍 Đối Soát Văn Bản Offline")
    c1, c2 = st.columns(2)
    with c1: f_a = st.file_uploader("Bản Gốc (A)", type=["pdf", "docx", "txt", "xlsx"], key="fa")
    with c2: f_b = st.file_uploader("Bản Mới (B)", type=["pdf", "docx", "txt", "xlsx"], key="fb")
    
    if st.button("🚀 BẮT ĐẦU SO SÁNH"):
        if f_a and f_b:
            t_a, t_b = read_file_content(f_a), read_file_content(f_b)
            diff = list(difflib.Differ().compare(t_a.splitlines(), t_b.splitlines()))
            for line in diff:
                if line.startswith('+ '): st.markdown(f"🟢 **Thêm:** `{line[2:]}`")
                elif line.startswith('- '): st.markdown(f"🔴 **Xóa:** ~~{line[2:]}~~")
        else: st.warning("Vui lòng tải đủ 2 bản!")

# --- HÀM TẠO FILE EXCEL MẪU ---
def tạo_excel_mẫu():
    # Danh sách các cột theo ảnh bạn gửi
    cột_mẫu = [
        "So", "Ten", "ChucVu", "Luong", "TenKhach", "TenSuKien", 
        "ThoiGian", "DiaDiem", "NgayCap", "LuongMoi", "LuongCu", 
        "NgayHieuLuc", "MaNV", "Phongban"
    ]
    
    # Dữ liệu mẫu ban đầu
    data_mẫu = [
        ["01", "Nguyễn Văn A", "Trưởng phòng", "20.000.000", "Lê Văn B", "Hội nghị khách hàng", 
         "08:00 01/02/2026", "Hà Nội", "01/01/2026", "25.000.000", "20.000.000", 
         "15/02/2026", "NV001", "Kinh doanh"]
    ]
    
    df_mẫu = pd.DataFrame(data_mẫu, columns=cột_mẫu)
    
    output = BytesIO()
    with pd.ExcelWriter(output, engine='xlsxwriter') as writer:
        df_mẫu.to_excel(writer, index=False, sheet_name='Mau_Nhap_Lieu')
        # Định dạng một chút cho đẹp
        workbook = writer.book
        header_fmt = workbook.add_format({'bold': True, 'bg_color': '#2ecc71', 'font_color': 'white', 'border': 1})
        for col_num, value in enumerate(df_mẫu.columns.values):
            writer.sheets['Mau_Nhap_Lieu'].write(0, col_num, value, header_fmt)
            writer.sheets['Mau_Nhap_Lieu'].set_column(col_num, col_num, 15)
            
    return output.getvalue()

# --- GIAO DIỆN TAB 3: TRỘN HỒ SƠ & HỢP ĐỒNG ---
with tabs[2]:
    st.header("🎭 Trộn Hồ Sơ Chuyên Nghiệp (V5)")
    
    # --- PHẦN 1: TẢI FILE ĐẦU VÀO ---
    st.subheader("📁 Bước 1: Tải dữ liệu và Mẫu Word")
    col_upload_ex, col_upload_wd = st.columns(2)
    
    with col_upload_ex:
        uploaded_excel = st.file_uploader("📂 Tải file Excel dữ liệu", type=["xlsx", "xls"])
        st.download_button("📥 Tải Excel mẫu (nếu chưa có)", tạo_excel_mẫu(), "Mau_Excel.xlsx")
        
    with col_upload_wd:
        uploaded_word = st.file_uploader("📂 Tải file Word mẫu ({{ }})", type=["docx"])
        st.download_button("📥 Tải Word mẫu (nếu chưa có)", tao_file_word_mau_giay_moi(), "Mau_Giay_Moi.docx")

    # Đọc dữ liệu từ Excel vào DataFrame nếu có file tải lên
    if uploaded_excel:
        st.session_state.df_merge = pd.read_excel(uploaded_excel)
    
    st.divider()

    # --- PHẦN 2: BẢNG CHỈNH SỬA DỮ LIỆU ---
    st.subheader("📝 Bước 2: Kiểm tra & Chỉnh sửa dữ liệu")
    st.write("*(Chữ đen đậm, bạn có thể sửa trực tiếp hoặc dán thêm dòng)*")
    
    # Cấu hình bảng chữ đen đậm
    config_cols = {col: st.column_config.TextColumn(label=f"**{col}**") for col in st.session_state.df_merge.columns}
    
    edited_df = st.data_editor(
        st.session_state.df_merge, 
        num_rows="dynamic", 
        use_container_width=True,
        column_config=config_cols,
        key="editor_v5"
    )

    st.divider()

    # --- PHẦN 3: LỰA CHỌN XUẤT FILE ---
    st.subheader("🚀 Bước 3: Lựa chọn kiểu xuất bản")
    
    if not edited_df.empty and uploaded_word:
        mode = st.radio("Chọn hình thức kết xuất:", 
                        ["📦 Xuất các file Word lẻ (Nén trong .ZIP)", 
                         "📄 Gộp tất cả vào 1 file Word duy nhất"])
        
        if st.button("🔥 BẮT ĐẦU XỬ LÝ"):
            all_docs = []
            try:
                # Tiến hành trộn dữ liệu
                for index, row in edited_df.iterrows():
                    doc = DocxTemplate(uploaded_word)
                    context = row.to_dict()
                    
                    # Logic đọc số tiền (nếu có)
                    if "LuongMoi" in context: context["LuongMoiChu"] = doc_so_thanh_chu_logic(str(context["LuongMoi"]))
                    
                    doc.render(context)
                    out_word = BytesIO()
                    doc.save(out_word)
                    all_docs.append({'data': out_word.getvalue(), 'name': str(row.get('Ten', f'File_{index+1}'))})

                # Trả kết quả theo lựa chọn
                if "lẻ" in mode:
                    zip_buffer = BytesIO()
                    with zipfile.ZipFile(zip_buffer, "a", zipfile.ZIP_DEFLATED, False) as zip_file:
                        for d in all_docs:
                            zip_file.writestr(f"{d['name'].replace(' ', '_')}.docx", d['data'])
                    st.success("✅ Đã tạo xong bộ file lẻ!")
                    st.download_button("📥 TẢI FILE .ZIP", zip_buffer.getvalue(), "Ket_Qua_Le.zip", use_container_width=True)
                
                else:
                    merged_data = gop_cac_file_word([d['data'] for d in all_docs])
                    st.success("✅ Đã gộp thành công vào 1 file duy nhất!")
                    st.download_button("📥 TẢI FILE WORD TỔNG", merged_data, "Ket_Qua_Tong_Hop.docx", use_container_width=True)
                    
            except Exception as e:
                st.error(f"❌ Lỗi: {e}")
    else:
        st.warning("⚠️ Vui lòng tải đủ file Excel và Word để bắt đầu.")

with tabs[3]: st.write("Chức năng đang phát triển...")
with tabs[4]: st.write("Chức năng đang phát triển...")
